from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

out = '/home/ubuntu/cong-nghe-12-nls-google-form/bao-cao-mau-nhom-1-khoang-5-diem.docx'
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DẤU VẾT TRI THỨC CÔNG NGHỆ 12')
r.bold = True
r.font.size = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO MẪU NHÓM 1 — MỨC THAM KHẢO KHOẢNG 5 ĐIỂM')
r.bold = True
r.font.size = Pt(14)

info = doc.add_table(rows=4, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = 'Table Grid'
rows = [
    ('Chủ đề', 'Ứng dụng AI trong phát hiện suy thoái rừng'),
    ('Nhóm', 'N1 — Giới thiệu chung về lâm nghiệp'),
    ('Lớp / mã lớp', '12A1 / 12A1'),
    ('Mã nhóm', 'N1'),
]
for row, (a, b) in zip(info.rows, rows):
    row.cells[0].text = a
    row.cells[1].text = b

doc.add_heading('1. Mở đầu', level=1)
doc.add_paragraph('Rừng có vai trò quan trọng đối với môi trường và đời sống. Hiện nay, công nghệ AI có thể hỗ trợ con người quan sát ảnh vệ tinh để phát hiện nơi có độ che phủ giảm. Nhóm chọn đề tài này vì muốn tìm hiểu cách công nghệ giúp theo dõi rừng nhanh hơn.')

doc.add_heading('2. AI có thể hỗ trợ những việc gì?', level=1)
doc.add_paragraph('AI có thể đọc dữ liệu ảnh vệ tinh hoặc ảnh chụp từ drone. Nếu ảnh ở hai thời điểm có sự khác nhau lớn, hệ thống có thể đánh dấu khu vực cần chú ý. AI cũng có thể hỗ trợ nhận diện cây bị sâu bệnh nếu có hình ảnh phù hợp. Kết quả của AI là cảnh báo hoặc dự đoán, không phải kết luận cuối cùng.')

doc.add_heading('3. Ví dụ minh họa đơn giản', level=1)
doc.add_paragraph('Nhóm giả sử một khu rừng được chụp ảnh vào tháng 1 và tháng 6. Ảnh tháng 6 cho thấy một phần màu xanh giảm so với ảnh trước. AI có thể đánh dấu khu vực này để kiểm tra. Tuy nhiên, nguyên nhân có thể là chặt cây, cháy rừng, mây che ảnh hoặc thay đổi mùa. Kiểm lâm cần đến hiện trường hoặc xem thêm dữ liệu trước khi quyết định.')

doc.add_paragraph('Minh chứng của nhóm hiện chỉ là sơ đồ quy trình tự vẽ và một ví dụ giả định, chưa có bộ ảnh vệ tinh thật hoặc số liệu đo cụ thể.')

doc.add_heading('4. Lợi ích và hạn chế', level=1)
doc.add_paragraph('Lợi ích là AI giúp xử lí nhiều ảnh trong thời gian ngắn và gợi ý khu vực cần ưu tiên. Hạn chế là kết quả phụ thuộc vào chất lượng ảnh và dữ liệu huấn luyện. Nếu dữ liệu không đủ, AI có thể nhận diện sai. Vì vậy con người vẫn phải kiểm tra, giải thích và chịu trách nhiệm về quyết định bảo vệ rừng.')

doc.add_heading('5. Kết luận', level=1)
doc.add_paragraph('AI có thể hỗ trợ phát hiện dấu hiệu suy thoái rừng nhưng không thay thế kiểm lâm. Nhóm hiểu rằng cần kết hợp dữ liệu, kiểm tra thực tế và kinh nghiệm chuyên môn. Báo cáo mới dừng ở mức giới thiệu, cần bổ sung ảnh có nguồn, số liệu và phân tích sâu hơn nếu có thêm thời gian.')

doc.add_heading('Tự đánh giá theo rubric', level=1)
t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'
for cell, text in zip(t.rows[0].cells, ['Tiêu chí', 'Mức chọn', 'Nhận xét ngắn']):
    cell.text = text
entries = [
    ('Nội dung kiến thức', '3/4', 'Có ý chính và phân biệt cảnh báo AI với quyết định con người, nhưng giải thích còn ngắn.'),
    ('Minh chứng thực hành', '2/4', 'Có ví dụ giả định và sơ đồ tự vẽ, chưa có ảnh/số liệu nguồn.'),
    ('Kỹ năng trình bày', '2/4', 'Bố cục dễ theo dõi nhưng lập luận chưa sâu.'),
    ('Slide / hình thức trực quan', '3/4', 'Có sơ đồ đơn giản, chưa có nhiều hình ảnh minh họa.'),
    ('Trả lời & phối hợp nhóm', '2/4', 'Có thể trả lời câu hỏi cơ bản, vai trò nhóm chưa thể hiện rõ.'),
]
for row in entries:
    cells = t.add_row().cells
    for cell, text in zip(cells, row):
        cell.text = text

doc.add_paragraph('Điểm tự chấm dự kiến: khoảng 48/100, tương đương khoảng 4,8/10. Đây là báo cáo mẫu cố ý ở mức trung bình để kiểm thử tính năng chấm điểm; không dùng làm bài nộp chính thức.')
doc.save(out)
print(out)
