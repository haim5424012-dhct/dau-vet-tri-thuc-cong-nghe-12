from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

out = '/home/ubuntu/cong-nghe-12-nls-google-form/bao-cao-mau-nhom-1-khoang-8-9-diem.docx'
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(12)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DẤU VẾT TRI THỨC CÔNG NGHỆ 12'); r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO MẪU NHÓM 1 — MỨC THAM KHẢO 8–9 ĐIỂM'); r.bold = True; r.font.size = Pt(14)

info = doc.add_table(rows=4, cols=2); info.alignment = WD_TABLE_ALIGNMENT.CENTER; info.style = 'Table Grid'
for row, (a,b) in zip(info.rows, [('Chủ đề','Theo dõi nhiệm vụ trồng và chăm sóc rừng bằng dữ liệu hiện trường'),('Nhóm','N1 — Vai trò và nhiệm vụ trồng rừng'),('Lớp / mã lớp','12A1 / 12A1'),('Mã nhóm','N1')]):
    row.cells[0].text = a; row.cells[1].text = b

doc.add_heading('1. Vấn đề và mục tiêu', level=1)
doc.add_paragraph('Suy thoái rừng thường diễn ra âm thầm qua mất cây, cháy rừng, sâu bệnh hoặc thay đổi sử dụng đất. Nhóm đề xuất tìm hiểu cách dùng AI để phân tích ảnh vệ tinh theo thời gian, từ đó phát hiện khu vực có dấu hiệu bất thường và hỗ trợ kiểm lâm ưu tiên kiểm tra. Mục tiêu của AI là tạo cảnh báo sớm, không tự đưa ra quyết định xử lí.')

doc.add_heading('2. Quy trình công nghệ đề xuất', level=1)
doc.add_paragraph('Quy trình gồm năm bước: thu thập ảnh vệ tinh ở hai hoặc nhiều thời điểm; tiền xử lí để giảm ảnh hưởng của mây và sai khác ánh sáng; trích xuất các chỉ số thực vật như NDVI; dùng mô hình AI để phát hiện vùng thay đổi; cuối cùng, kiểm lâm đối chiếu với bản đồ, hồ sơ và kiểm tra thực địa. Nếu dữ liệu chưa đủ, kết quả phải được ghi là cảnh báo sơ bộ.')

doc.add_heading('3. Minh chứng và diễn giải', level=1)
doc.add_paragraph('Trong mô hình minh họa, ảnh tháng 1 và tháng 6 được đặt cạnh nhau. Khu vực có chỉ số xanh giảm liên tục được tô màu vàng để ưu tiên kiểm tra. Nhóm xây dựng sơ đồ dữ liệu: ảnh vệ tinh → tiền xử lí → AI phát hiện thay đổi → kiểm tra thực địa → quyết định của con người. Ví dụ cho thấy AI giúp rút ngắn bước sàng lọc, nhưng nguyên nhân suy giảm có thể do chặt cây, cháy, mây che ảnh hoặc mùa khô; vì vậy không thể kết luận chỉ từ màu trên ảnh.')

doc.add_heading('4. Lợi ích, rủi ro và trách nhiệm', level=1)
doc.add_paragraph('Lợi ích chính là theo dõi diện rộng, phát hiện thay đổi nhanh và lưu được lịch sử dữ liệu. Rủi ro gồm ảnh thiếu chất lượng, dữ liệu huấn luyện thiên lệch, cảnh báo nhầm và bỏ sót khu vực nhỏ. Người sử dụng cần kiểm tra chất lượng ảnh, ghi nguồn dữ liệu, giữ nhật ký quyết định và không dùng cảnh báo AI làm căn cứ duy nhất để xử phạt. Đây là điểm quan trọng của việc kết hợp năng lực số với trách nhiệm nghề nghiệp.')

doc.add_heading('5. Kết luận và hướng phát triển', level=1)
doc.add_paragraph('AI phù hợp với vai trò trợ lí phân tích trong giám sát rừng. Hướng phát triển tiếp theo là bổ sung ảnh có nguồn rõ ràng, gắn tọa độ, so sánh với số liệu thực địa và đánh giá độ chính xác bằng các mẫu đã được kiểm chứng. Nhóm nhận thấy giải pháp có giá trị khi được đặt trong quy trình phối hợp giữa dữ liệu, kiểm lâm và cộng đồng địa phương.')

doc.add_heading('Tự đánh giá theo rubric', level=1)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
for cell, text in zip(t.rows[0].cells, ['Tiêu chí','Mức chọn','Nhận xét']): cell.text = text
for row in [
 ('Nội dung kiến thức','4/4','Giải thích đúng vai trò AI, quy trình và trách nhiệm con người.'),
 ('Minh chứng thực hành','3/4','Có sơ đồ và ví dụ diễn giải, nhưng chưa có dữ liệu vệ tinh thật.'),
 ('Kỹ năng trình bày','3/4','Lập luận có trình tự, nêu được rủi ro và hướng phát triển.'),
 ('Slide / hình thức trực quan','4/4','Có sơ đồ quy trình và cặp ảnh minh họa theo bố cục rõ ràng.'),
 ('Trả lời & phối hợp nhóm','3/4','Có thể trả lời câu hỏi chính, cần mô tả vai trò từng thành viên rõ hơn.'),
]:
    cells=t.add_row().cells
    for cell,text in zip(cells,row): cell.text=text

doc.add_paragraph('Điểm tự chấm dự kiến: 86/100, tương đương khoảng 8,6/10. Đây là báo cáo mẫu dùng để kiểm thử mức điểm cao hơn, không phải bài nộp chính thức.')
doc.save(out)
print(out)
