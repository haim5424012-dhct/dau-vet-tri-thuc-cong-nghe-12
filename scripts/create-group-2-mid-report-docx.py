from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

out = "/home/ubuntu/cong-nghe-12-nls-google-form/bao-cao-mau-nhom-2-khoang-6-7-diem.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Pt(45); section.bottom_margin = Pt(45)
styles = doc.styles
styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(11)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DẤU VẾT TRI THỨC CÔNG NGHỆ 12"); r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BÁO CÁO MẪU NHÓM 2 — TRỒNG VÀ CHĂM SÓC RỪNG"); r.bold = True; r.font.size = Pt(14)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Mức tham khảo: khoảng 6–7 điểm").italic = True

doc.add_heading("1. Thông tin chung", level=1)
for label, value in [("Nhóm", "N2 — Trồng và chăm sóc rừng"), ("Chủ đề", "Ứng dụng ảnh vệ tinh và drone theo dõi độ che phủ rừng"), ("Mã lớp", "12A1"), ("Mã nhóm", "N2")]:
    doc.add_paragraph(f"{label}: {value}")

doc.add_heading("2. Nội dung tìm hiểu", level=1)
doc.add_paragraph("Ảnh vệ tinh có thể giúp quan sát diện tích và độ che phủ của rừng theo thời gian. Khi so sánh các ảnh ở những thời điểm khác nhau, người sử dụng có thể nhận ra một số khu vực cây bị thưa hoặc đất trống tăng lên. Drone có thể bay thấp để chụp ảnh chi tiết hơn ở một khu vực nhỏ.")
doc.add_paragraph("Theo nhóm, dữ liệu ảnh giúp cán bộ lâm nghiệp biết nơi cần kiểm tra trước. Một số phần mềm có thể hỗ trợ phân loại ảnh và ước lượng diện tích tán cây. Tuy nhiên, kết quả của phần mềm vẫn có thể nhầm giữa cây non, bụi cây và đất có màu tương tự.")

doc.add_heading("3. Đề xuất giải pháp", level=1)
doc.add_paragraph("Nhóm đề xuất dùng ảnh vệ tinh để theo dõi theo tháng, kết hợp drone khi phát hiện khu vực có thay đổi lớn. Cán bộ sẽ kiểm tra thực địa một số điểm, ghi lại nguyên nhân như khai thác, cháy rừng, sâu bệnh hoặc thời tiết. Sau đó mới quyết định trồng bổ sung hay chăm sóc.")

doc.add_heading("4. Hạn chế và trách nhiệm", level=1)
doc.add_paragraph("Báo cáo chưa có số liệu đo thực tế và chưa so sánh nhiều loại ảnh. Nhóm cũng chưa thử nghiệm một phần mềm cụ thể. Vì vậy, cảnh báo của công nghệ chỉ nên được xem là thông tin hỗ trợ; quyết định cuối cùng cần có người có chuyên môn kiểm tra và chịu trách nhiệm.")

doc.add_heading("5. Kết luận", level=1)
doc.add_paragraph("Ảnh vệ tinh và drone có thể làm cho việc theo dõi rừng nhanh hơn, nhất là ở khu vực rộng. Muốn sử dụng hiệu quả cần kết hợp dữ liệu số với kiểm tra thực địa và ghi chép trung thực.")

doc.add_heading("Tự đánh giá rubric dự kiến", level=1)
for item in ["Nội dung kiến thức: 3/4", "Minh chứng thực hành: 2/4", "Kỹ năng trình bày: 3/4", "Slide / hình thức trực quan: 2/4", "Trả lời & phối hợp nhóm: 2/4"]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph("Điểm quy đổi dự kiến trên website: 63/100, tương đương 6,3/10.")

doc.save(out)
print(out)
